"""
Document processing module for extracting text from various file formats
"""

import os
import re
import asyncio
from typing import List, Dict, Any, Optional
import magic
from datetime import datetime

# Document processing libraries
import PyPDF2
import docx
from bs4 import BeautifulSoup
import email
from email.mime.text import MIMEText

from ..models.document import DocumentMetadata, DocumentType
from ..utils.logger import get_logger

logger = get_logger(__name__)

class DocumentProcessor:
    """Handles document ingestion and processing"""
    
    def __init__(self):
        self.supported_formats = {
            DocumentType.PDF: self._process_pdf,
            DocumentType.DOCX: self._process_docx,
            DocumentType.TXT: self._process_text,
            DocumentType.HTML: self._process_html,
            DocumentType.EMAIL: self._process_email
        }
        
        # Text cleaning patterns
        self.cleaning_patterns = [
            (r'\s+', ' '),  # Multiple whitespace to single space
            (r'\n\s*\n', '\n\n'),  # Multiple newlines to double newline
            (r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\<\>\~\`\|\\]', ''),  # Remove special chars
        ]
        
    async def process_document(self, file_path: str, doc_metadata: DocumentMetadata) -> List[str]:
        """Process document and extract text chunks"""
        try:
            logger.info(f"Processing document: {doc_metadata.filename} ({doc_metadata.doc_type})")
            
            # Verify file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Get appropriate processor
            processor = self.supported_formats.get(doc_metadata.doc_type)
            if not processor:
                raise ValueError(f"Unsupported document type: {doc_metadata.doc_type}")
            
            # Process document
            raw_text = await processor(file_path)
            
            # Clean and chunk text
            cleaned_text = self._clean_text(raw_text)
            chunks = self._split_into_chunks(cleaned_text, max_chunk_size=500)
            
            # Filter empty chunks
            chunks = [chunk for chunk in chunks if chunk.strip()]
            
            logger.info(f"Extracted {len(chunks)} chunks from {doc_metadata.filename}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing document {doc_metadata.doc_id}: {str(e)}")
            raise
    
    async def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            # Add page marker for reference
                            text_content.append(f"[PAGE {page_num + 1}]\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            raise
    
    async def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
    
    async def _process_html(self, file_path: str) -> str:
        """Extract text from HTML"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {str(e)}")
            raise
    
    async def _process_email(self, file_path: str) -> str:
        """Extract text from email"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                msg = email.message_from_file(file)
            
            text_parts = []
            
            # Extract headers
            headers = [
                f"From: {msg.get('From', 'Unknown')}",
                f"To: {msg.get('To', 'Unknown')}",
                f"Subject: {msg.get('Subject', 'No Subject')}",
                f"Date: {msg.get('Date', 'Unknown')}"
            ]
            text_parts.append('\n'.join(headers))
            text_parts.append('-' * 50)
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            text_parts.append(payload.decode('utf-8', errors='ignore'))
                    elif part.get_content_type() == "text/html":
                        payload = part.get_payload(decode=True)
                        if payload:
                            soup = BeautifulSoup(payload.decode('utf-8', errors='ignore'), 'html.parser')
                            text_parts.append(soup.get_text())
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    text_parts.append(payload.decode('utf-8', errors='ignore'))
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error processing email {file_path}: {str(e)}")
            raise
    
    async def _process_text(self, file_path: str) -> str:
        """Process plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as file:
                content = file.read()
                return content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Apply cleaning patterns
        cleaned_text = text
        for pattern, replacement in self.cleaning_patterns:
            cleaned_text = re.sub(pattern, replacement, cleaned_text)
        
        # Remove excessive whitespace
        cleaned_text = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned_text)
        cleaned_text = re.sub(r'[ \t]+', ' ', cleaned_text)
        
        return cleaned_text.strip()
    
    def _split_into_chunks(self, text: str, max_chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into semantic chunks with overlap"""
        if not text:
            return []
        
        # First, split by double newlines (paragraphs)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            # If paragraph is too long, split by sentences
            if len(paragraph) > max_chunk_size:
                sentences = self._split_by_sentences(paragraph)
                
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) + 2 <= max_chunk_size:
                        current_chunk += sentence + ". "
                    else:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = sentence + ". "
            else:
                # Check if adding this paragraph exceeds chunk size
                if len(current_chunk) + len(paragraph) + 2 <= max_chunk_size:
                    current_chunk += paragraph + "\n\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = paragraph + "\n\n"
        
        # Add remaining chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Add overlap between chunks
        if overlap > 0 and len(chunks) > 1:
            chunks = self._add_overlap(chunks, overlap)
        
        return chunks
    
    def _split_by_sentences(self, text: str) -> List[str]:
        """Split text by sentences"""
        # Simple sentence splitting - could be improved with spaCy or NLTK
        sentence_endings = r'[.!?]+(?:\s|$)'
        sentences = re.split(sentence_endings, text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _add_overlap(self, chunks: List[str], overlap: int) -> List[str]:
        """Add overlap between consecutive chunks"""
        if len(chunks) <= 1:
            return chunks
        
        overlapped_chunks = [chunks[0]]
        
        for i in range(1, len(chunks)):
            # Get last `overlap` characters from previous chunk
            prev_chunk = chunks[i-1]
            current_chunk = chunks[i]
            
            if len(prev_chunk) > overlap:
                overlap_text = prev_chunk[-overlap:]
                # Find word boundary
                space_idx = overlap_text.find(' ')
                if space_idx != -1:
                    overlap_text = overlap_text[space_idx+1:]
                
                overlapped_chunk = overlap_text + " " + current_chunk
            else:
                overlapped_chunk = current_chunk
            
            overlapped_chunks.append(overlapped_chunk)
        
        return overlapped_chunks
    
    def get_document_stats(self, text: str) -> Dict[str, Any]:
        """Get document statistics"""
        if not text:
            return {"word_count": 0, "char_count": 0, "paragraph_count": 0}
        
        words = text.split()
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        
        return {
            "word_count": len(words),
            "char_count": len(text),
            "paragraph_count": len(paragraphs),
            "estimated_reading_time": len(words) // 200  # Assuming 200 WPM
        }
    
    def detect_language(self, text: str) -> str:
        """Simple language detection (could be improved)"""
        # This is a very basic implementation
        # In production, you might want to use langdetect or similar
        
        if not text:
            return "unknown"
        
        # Simple heuristics for common languages
        english_words = ["the", "and", "or", "but", "in", "on", "at", "to", "for", "of", "with", "by"]
        sample_text = text.lower()[:1000]  # Use first 1000 chars
        
        english_count = sum(1 for word in english_words if word in sample_text)
        
        if english_count >= 3:
            return "en"
        else:
            return "unknown"
